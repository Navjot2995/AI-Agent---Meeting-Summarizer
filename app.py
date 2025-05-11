import streamlit as st
import os
from dotenv import load_dotenv
import groq
import speech_recognition as sr
from pydub import AudioSegment
import tempfile
import magic
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime
from docx import Document
from fpdf import FPDF
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.sentiment import SentimentIntensityAnalyzer
from nltk.corpus import stopwords
import json
import noisereduce as nr
import numpy as np
from scipy.io import wavfile
from moviepy.editor import VideoFileClip
import google.oauth2.credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
import pickle
import base64
import requests
from urllib.parse import urlparse
import mimetypes
from collections import Counter
import re
from textblob import TextBlob
import spacy
import networkx as nx
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import sys
import subprocess

# Configure page
st.set_page_config(
    page_title="Meeting Summarizer AI",
    page_icon="🎙️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')
try:
    nltk.data.find('sentiment/vader_lexicon')
except LookupError:
    nltk.download('vader_lexicon')
try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

# Load spaCy model with better error handling
def load_spacy_model():
    """Load spaCy model with proper error handling."""
    try:
        return spacy.load('en_core_web_sm')
    except OSError:
        try:
            # Try downloading the model if not found
            st.info("Downloading spaCy model... This may take a few minutes.")
            subprocess.run([sys.executable, "-m", "spacy", "download", "en_core_web_sm"], check=True)
            st.success("spaCy model downloaded successfully!")
            return spacy.load('en_core_web_sm')
        except Exception as e:
            st.error(f"""
            Failed to load spaCy model. Please ensure the model is installed:
            
            ```bash
            python -m spacy download en_core_web_sm
            ```
            
            Error details: {str(e)}
            """)
            st.stop()
    except Exception as e:
        st.error(f"Error loading spaCy model: {str(e)}")
        st.stop()

# Initialize spaCy with fallback
try:
    nlp = load_spacy_model()
except Exception as e:
    st.warning(f"""
    Failed to initialize spaCy. The application will continue with limited functionality.
    Error: {str(e)}
    """)
    # Create a dummy NLP object for basic functionality
    class DummyNLP:
        def __call__(self, text):
            return type('Doc', (), {'ents': []})()
        def pipe(self, texts):
            return [self(text) for text in texts]
    nlp = DummyNLP()

# Load environment variables
load_dotenv()

# Configure Groq
def get_groq_client():
    """Initialize and return Groq client with proper error handling."""
    try:
        # First try to get from Streamlit secrets
        if "GROQ_API_KEY" in st.secrets:
            api_key = st.secrets["GROQ_API_KEY"]
        # Then try environment variable
        elif "GROQ_API_KEY" in os.environ:
            api_key = os.environ["GROQ_API_KEY"]
        else:
            st.error("Groq API key not found. Please set it in Streamlit secrets or environment variables.")
            return None

        if not api_key:
            st.error("Groq API key is empty. Please provide a valid API key.")
            return None

        # Initialize Groq client with the correct method
        try:
            # Create client with minimal configuration
            client = groq.Client(api_key=api_key)
            
            # Test the client with a simple request
            response = client.chat.completions.create(
                model="mixtral-8x7b-32768",
                messages=[{"role": "user", "content": "test"}],
                max_tokens=5
            )
            return client
        except Exception as e:
            st.error(f"Error testing Groq client: {str(e)}")
            return None

    except Exception as e:
        st.error(f"Error initializing Groq client: {str(e)}")
        return None

# Initialize Groq client
groq_client = get_groq_client()

# Google Calendar API setup
SCOPES = ['https://www.googleapis.com/auth/calendar']

def get_google_calendar_service():
    """Set up Google Calendar API service."""
    creds = None
    if os.path.exists('token.pickle'):
        with open('token.pickle', 'rb') as token:
            creds = pickle.load(token)
    
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(
                'credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)
        with open('token.pickle', 'wb') as token:
            pickle.dump(creds, token)
    
    return build('calendar', 'v3', credentials=creds)

def extract_audio_from_video(video_file):
    """Extract audio from video file."""
    with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as temp_audio:
        temp_audio_path = temp_audio.name
    
    video = VideoFileClip(video_file)
    video.audio.write_audiofile(temp_audio_path)
    return temp_audio_path

def reduce_noise(audio_file):
    """Reduce background noise from audio file."""
    # Read audio file
    rate, data = wavfile.read(audio_file)
    
    # Perform noise reduction
    reduced_noise = nr.reduce_noise(y=data, sr=rate)
    
    # Save processed audio
    with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as temp_file:
        temp_file_path = temp_file.name
    wavfile.write(temp_file_path, rate, reduced_noise)
    
    return temp_file_path

def analyze_speaking_time(transcription, speaker_segments):
    """Analyze speaking time per speaker."""
    speaking_time = {}
    for speaker, segments in speaker_segments.items():
        total_time = sum(end - start for start, end in segments)
        speaking_time[speaker] = total_time
    
    return speaking_time

def create_analytics_dashboard(transcription, sentiments, speaking_time, action_items):
    """Create comprehensive analytics dashboard."""
    # Create tabs for different analytics
    tab1, tab2, tab3, tab4 = st.tabs(["Overview", "Sentiment Analysis", "Speaking Time", "Action Items"])
    
    with tab1:
        st.subheader("Meeting Overview")
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("Meeting Duration", f"{sum(speaking_time.values()):.2f} seconds")
        with col2:
            st.metric("Number of Speakers", len(speaking_time))
        with col3:
            st.metric("Action Items", len(action_items.split('\n')))
    
    with tab2:
        st.subheader("Sentiment Analysis")
        df = pd.DataFrame(sentiments)
        fig = px.line(df, y=['positive', 'negative', 'neutral'], 
                     title='Sentiment Analysis Over Time')
        st.plotly_chart(fig)
        
        # Sentiment distribution pie chart
        avg_sentiment = df[['positive', 'negative', 'neutral']].mean()
        fig = px.pie(values=avg_sentiment.values, 
                    names=avg_sentiment.index,
                    title='Overall Sentiment Distribution')
        st.plotly_chart(fig)
    
    with tab3:
        st.subheader("Speaking Time Analysis")
        # Speaking time bar chart
        fig = px.bar(x=list(speaking_time.keys()),
                    y=list(speaking_time.values()),
                    title='Speaking Time per Speaker')
        st.plotly_chart(fig)
    
    with tab4:
        st.subheader("Action Items Analysis")
        # Action items timeline
        action_items_list = action_items.split('\n')
        fig = go.Figure(data=[go.Table(
            header=dict(values=['Action Item', 'Assignee', 'Deadline'],
                       fill_color='paleturquoise',
                       align='left'),
            cells=dict(values=[[item.split(' - ')[0] for item in action_items_list],
                             [item.split(' - ')[1] if ' - ' in item else 'Unassigned' for item in action_items_list],
                             [item.split(' - ')[2] if ' - ' in item else 'No deadline' for item in action_items_list]],
                      fill_color='lavender',
                      align='left'))
        ])
        st.plotly_chart(fig)

def create_calendar_event(summary, start_time, end_time, description):
    """Create a Google Calendar event for the meeting."""
    try:
        service = get_google_calendar_service()
        event = {
            'summary': summary,
            'description': description,
            'start': {
                'dateTime': start_time.isoformat(),
                'timeZone': 'UTC',
            },
            'end': {
                'dateTime': end_time.isoformat(),
                'timeZone': 'UTC',
            },
        }
        event = service.events().insert(calendarId='primary', body=event).execute()
        return event.get('htmlLink')
    except Exception as e:
        st.error(f"Error creating calendar event: {str(e)}")
        return None

def convert_audio_to_wav(audio_file):
    """Convert uploaded audio file to WAV format."""
    with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as temp_wav:
        temp_wav_path = temp_wav.name
    
    # Convert to WAV using pydub
    audio = AudioSegment.from_file(audio_file)
    audio.export(temp_wav_path, format="wav")
    return temp_wav_path

def transcribe_audio(audio_file):
    """Transcribe audio file to text using SpeechRecognition."""
    recognizer = sr.Recognizer()
    
    with sr.AudioFile(audio_file) as source:
        audio_data = recognizer.record(source)
        try:
            text = recognizer.recognize_google(audio_data)
            return text
        except Exception as e:
            st.error(f"Error during transcription: {str(e)}")
            return None

def analyze_sentiment(text):
    """Analyze sentiment of the meeting transcript."""
    sia = SentimentIntensityAnalyzer()
    sentences = sent_tokenize(text)
    sentiments = []
    
    for sentence in sentences:
        sentiment = sia.polarity_scores(sentence)
        sentiments.append({
            'sentence': sentence,
            'compound': sentiment['compound'],
            'positive': sentiment['pos'],
            'negative': sentiment['neg'],
            'neutral': sentiment['neu']
        })
    
    return sentiments

def extract_action_items(text):
    """Extract action items from the transcript using Groq."""
    if not groq_client:
        st.error("Groq client not initialized. Please check your API key configuration.")
        return None

    try:
        response = groq_client.chat.completions.create(
            model="mixtral-8x7b-32768",
            messages=[
                {"role": "system", "content": "Extract action items from the following meeting transcript. Format them as a list with assignees and deadlines if mentioned."},
                {"role": "user", "content": text}
            ],
            max_tokens=300,
            temperature=0.7,
            top_p=0.95
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Error extracting action items: {str(e)}")
        return None

def summarize_text(text):
    """Summarize text using Groq's LLM."""
    if not groq_client:
        st.error("Groq client not initialized. Please check your API key configuration.")
        return None

    try:
        response = groq_client.chat.completions.create(
            model="mixtral-8x7b-32768",
            messages=[
                {"role": "system", "content": """You are a professional meeting summarizer. Create a comprehensive summary of the following meeting transcript. 
                Include:
                1. Key points discussed
                2. Decisions made
                3. Action items with assignees
                4. Timeline of important events
                5. Risk factors or concerns raised
                Format the summary in clear sections."""},
                {"role": "user", "content": text}
            ],
            max_tokens=1000,
            temperature=0.7,
            top_p=0.95
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Error during summarization: {str(e)}")
        return None

def export_to_pdf(summary, action_items, sentiment_analysis):
    """Export meeting summary to PDF."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(200, 10, "Meeting Summary", ln=True, align="C")
    pdf.set_font("Arial", "", 12)
    
    # Add summary
    pdf.ln(10)
    pdf.set_font("Arial", "B", 14)
    pdf.cell(200, 10, "Summary", ln=True)
    pdf.set_font("Arial", "", 12)
    pdf.multi_cell(0, 10, summary)
    
    # Add action items
    pdf.ln(10)
    pdf.set_font("Arial", "B", 14)
    pdf.cell(200, 10, "Action Items", ln=True)
    pdf.set_font("Arial", "", 12)
    pdf.multi_cell(0, 10, action_items)
    
    # Save PDF
    pdf_path = "meeting_summary.pdf"
    pdf.output(pdf_path)
    return pdf_path

def export_to_docx(summary, action_items, sentiment_analysis):
    """Export meeting summary to DOCX."""
    doc = Document()
    doc.add_heading("Meeting Summary", 0)
    
    # Add summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary)
    
    # Add action items
    doc.add_heading("Action Items", level=1)
    doc.add_paragraph(action_items)
    
    # Save DOCX
    docx_path = "meeting_summary.docx"
    doc.save(docx_path)
    return docx_path

def download_meeting_recording(url):
    """Download meeting recording from various platforms with enhanced error handling."""
    try:
        # Parse the URL to determine the platform
        parsed_url = urlparse(url)
        domain = parsed_url.netloc.lower()

        # Handle different meeting platforms
        if 'zoom.us' in domain:
            # Zoom recording download
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            response = requests.get(url, stream=True, headers=headers, timeout=30)
            if response.status_code == 200:
                return response.content
            elif response.status_code == 403:
                st.error("Access denied. Please check if the Zoom recording is publicly accessible.")
                return None
            elif response.status_code == 404:
                st.error("Recording not found. Please verify the Zoom recording URL.")
                return None
            else:
                st.error(f"Failed to download Zoom recording. Status code: {response.status_code}")
                return None
        elif 'teams.microsoft.com' in domain:
            # Teams recording download
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            response = requests.get(url, stream=True, headers=headers, timeout=30)
            if response.status_code == 200:
                return response.content
            elif response.status_code == 401:
                st.error("Authentication required. Please ensure you're logged into Microsoft Teams.")
                return None
            elif response.status_code == 403:
                st.error("Access denied. Please check your Teams permissions.")
                return None
            else:
                st.error(f"Failed to download Teams recording. Status code: {response.status_code}")
                return None
        elif 'meet.google.com' in domain:
            # Google Meet recording download
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            response = requests.get(url, stream=True, headers=headers, timeout=30)
            if response.status_code == 200:
                return response.content
            elif response.status_code == 401:
                st.error("Authentication required. Please ensure you're logged into Google.")
                return None
            elif response.status_code == 403:
                st.error("Access denied. Please check your Google Drive permissions.")
                return None
            else:
                st.error(f"Failed to download Google Meet recording. Status code: {response.status_code}")
                return None
        else:
            st.error("Unsupported meeting platform. Please provide a Zoom, Teams, or Google Meet recording URL.")
            return None
    except requests.exceptions.Timeout:
        st.error("Download timed out. Please check your internet connection and try again.")
        return None
    except requests.exceptions.ConnectionError:
        st.error("Connection error. Please check your internet connection and try again.")
        return None
    except Exception as e:
        st.error(f"Error downloading meeting recording: {str(e)}")
        return None

def extract_key_topics(text):
    """Extract key topics from the meeting transcript."""
    try:
        # Tokenize and remove stopwords
        stop_words = set(stopwords.words('english'))
        words = word_tokenize(text.lower())
        words = [word for word in words if word.isalnum() and word not in stop_words]
        
        # Get word frequencies
        word_freq = Counter(words)
        
        # Extract named entities using spaCy
        doc = nlp(text)
        entities = [ent.text for ent in doc.ents if ent.label_ in ['ORG', 'PRODUCT', 'PERSON', 'GPE']]
        
        # Combine word frequencies and entities
        topics = word_freq.most_common(10) + [(entity, 1) for entity in set(entities)]
        
        return topics
    except Exception as e:
        st.error(f"Error extracting key topics: {str(e)}")
        return []

def analyze_speaker_engagement(transcription, speaker_segments):
    """Analyze speaker engagement and interaction patterns."""
    try:
        # Create a graph for speaker interactions
        G = nx.Graph()
        
        # Add nodes for each speaker
        for speaker in speaker_segments.keys():
            G.add_node(speaker)
        
        # Analyze interactions
        interactions = []
        for i, (speaker1, segments1) in enumerate(speaker_segments.items()):
            for speaker2, segments2 in list(speaker_segments.items())[i+1:]:
                # Count interactions between speakers
                interaction_count = sum(1 for s1 in segments1 for s2 in segments2 
                                     if abs(s1[0] - s2[0]) < 5)  # 5-second threshold
                if interaction_count > 0:
                    G.add_edge(speaker1, speaker2, weight=interaction_count)
                    interactions.append((speaker1, speaker2, interaction_count))
        
        return G, interactions
    except Exception as e:
        st.error(f"Error analyzing speaker engagement: {str(e)}")
        return None, []

def analyze_meeting_structure(transcription):
    """Analyze meeting structure and flow."""
    try:
        # Split into sentences
        sentences = sent_tokenize(transcription)
        
        # Analyze sentence types
        sentence_types = {
            'questions': [],
            'statements': [],
            'decisions': []
        }
        
        for sentence in sentences:
            if '?' in sentence:
                sentence_types['questions'].append(sentence)
            elif any(word in sentence.lower() for word in ['decide', 'decided', 'agreed', 'agreement']):
                sentence_types['decisions'].append(sentence)
            else:
                sentence_types['statements'].append(sentence)
        
        return sentence_types
    except Exception as e:
        st.error(f"Error analyzing meeting structure: {str(e)}")
        return {}

def create_enhanced_analytics_dashboard(transcription, sentiments, speaking_time, action_items, topics, speaker_engagement, meeting_structure):
    """Create an enhanced analytics dashboard with more insights."""
    # Create tabs for different analytics
    tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
        "Overview", "Sentiment Analysis", "Speaking Time", "Action Items",
        "Key Topics", "Speaker Engagement", "Meeting Structure"
    ])
    
    with tab1:
        st.subheader("Meeting Overview")
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Meeting Duration", f"{sum(speaking_time.values()):.2f} seconds")
        with col2:
            st.metric("Number of Speakers", len(speaking_time))
        with col3:
            st.metric("Action Items", len(action_items.split('\n')))
        with col4:
            st.metric("Key Topics", len(topics))
    
    with tab2:
        st.subheader("Sentiment Analysis")
        df = pd.DataFrame(sentiments)
        fig = px.line(df, y=['positive', 'negative', 'neutral'], 
                     title='Sentiment Analysis Over Time')
        st.plotly_chart(fig)
        
        # Sentiment distribution pie chart
        avg_sentiment = df[['positive', 'negative', 'neutral']].mean()
        fig = px.pie(values=avg_sentiment.values, 
                    names=avg_sentiment.index,
                    title='Overall Sentiment Distribution')
        st.plotly_chart(fig)
    
    with tab3:
        st.subheader("Speaking Time Analysis")
        # Speaking time bar chart
        fig = px.bar(x=list(speaking_time.keys()),
                    y=list(speaking_time.values()),
                    title='Speaking Time per Speaker')
        st.plotly_chart(fig)
    
    with tab4:
        st.subheader("Action Items Analysis")
        # Action items timeline
        action_items_list = action_items.split('\n')
        fig = go.Figure(data=[go.Table(
            header=dict(values=['Action Item', 'Assignee', 'Deadline'],
                       fill_color='paleturquoise',
                       align='left'),
            cells=dict(values=[[item.split(' - ')[0] for item in action_items_list],
                             [item.split(' - ')[1] if ' - ' in item else 'Unassigned' for item in action_items_list],
                             [item.split(' - ')[2] if ' - ' in item else 'No deadline' for item in action_items_list]],
                      fill_color='lavender',
                      align='left'))
        ])
        st.plotly_chart(fig)
    
    with tab5:
        st.subheader("Key Topics Analysis")
        # Display key topics
        topics_df = pd.DataFrame(topics, columns=['Topic', 'Frequency'])
        fig = px.bar(topics_df, x='Topic', y='Frequency',
                    title='Key Topics Frequency')
        st.plotly_chart(fig)
    
    with tab6:
        st.subheader("Speaker Engagement Analysis")
        if speaker_engagement[0]:
            G, interactions = speaker_engagement
            # Create network graph
            pos = nx.spring_layout(G)
            fig = go.Figure()
            
            # Add edges
            for edge in G.edges(data=True):
                x0, y0 = pos[edge[0]]
                x1, y1 = pos[edge[1]]
                fig.add_trace(go.Scatter(
                    x=[x0, x1, None], y=[y0, y1, None],
                    mode='lines',
                    line=dict(width=edge[2]['weight']),
                    hoverinfo='none'
                ))
            
            # Add nodes
            for node in G.nodes():
                x, y = pos[node]
                fig.add_trace(go.Scatter(
                    x=[x], y=[y],
                    mode='markers+text',
                    marker=dict(size=20),
                    text=[node],
                    hoverinfo='text'
                ))
            
            fig.update_layout(showlegend=False)
            st.plotly_chart(fig)
    
    with tab7:
        st.subheader("Meeting Structure Analysis")
        # Display meeting structure statistics
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Questions Asked", len(meeting_structure.get('questions', [])))
        with col2:
            st.metric("Decisions Made", len(meeting_structure.get('decisions', [])))
        with col3:
            st.metric("Statements Made", len(meeting_structure.get('statements', [])))
        
        # Display key decisions
        if meeting_structure.get('decisions'):
            st.write("Key Decisions:")
            for decision in meeting_structure['decisions']:
                st.write(f"- {decision}")

def process_meeting_file(temp_file_path):
    """Process the meeting file and generate enhanced analysis."""
    try:
        # Handle video files
        if temp_file_path.lower().endswith(('.mp4', '.avi', '.mov')):
            temp_file_path = extract_audio_from_video(temp_file_path)

        # Convert to WAV if necessary
        if not temp_file_path.lower().endswith('.wav'):
            temp_file_path = convert_audio_to_wav(temp_file_path)

        # Reduce noise
        temp_file_path = reduce_noise(temp_file_path)

        # Transcribe audio
        st.write("Transcribing audio...")
        transcription = transcribe_audio(temp_file_path)
        
        if transcription:
            # Create tabs for different sections
            tab1, tab2, tab3, tab4, tab5 = st.tabs([
                "Transcription", "Summary", "Action Items", "Analysis", "Enhanced Analytics"
            ])
            
            with tab1:
                st.subheader("Transcription")
                st.write(transcription)

            # Generate enhanced summary
            st.write("Generating comprehensive summary...")
            summary = summarize_text(transcription)
            
            with tab2:
                if summary:
                    st.subheader("Summary")
                    st.write(summary)

            # Extract action items
            action_items = extract_action_items(transcription)
            
            with tab3:
                if action_items:
                    st.subheader("Action Items")
                    st.write(action_items)

            # Analyze sentiment
            sentiments = analyze_sentiment(transcription)
            
            with tab4:
                st.subheader("Sentiment Analysis")
                df = pd.DataFrame(sentiments)
                fig = px.line(df, y=['positive', 'negative', 'neutral'], 
                            title='Sentiment Analysis Over Time')
                st.plotly_chart(fig)

            # Enhanced analysis
            st.write("Performing enhanced analysis...")
            topics = extract_key_topics(transcription)
            speaker_engagement = analyze_speaker_engagement(transcription, 
                {"Speaker 1": [(0, 100), (200, 300)], "Speaker 2": [(100, 200), (300, 400)]})
            meeting_structure = analyze_meeting_structure(transcription)

            # Create enhanced analytics dashboard
            with tab5:
                create_enhanced_analytics_dashboard(
                    transcription,
                    sentiments,
                    {"Speaker 1": [(0, 100), (200, 300)], "Speaker 2": [(100, 200), (300, 400)]},
                    action_items,
                    topics,
                    speaker_engagement,
                    meeting_structure
                )

    except Exception as e:
        st.error(f"Error processing meeting file: {str(e)}")
    finally:
        # Clean up temporary files
        try:
            os.unlink(temp_file_path)
        except:
            pass

def main():
    st.title("🎙️ Advanced Meeting Summarizer AI")
    st.write("Upload your meeting recording or provide a meeting recording URL for comprehensive analysis.")

    # Sidebar for additional options
    with st.sidebar:
        st.header("Options")
        export_format = st.selectbox(
            "Choose export format",
            ["PDF", "DOCX", "TXT"]
        )
        
        # Calendar integration option
        create_calendar = st.checkbox("Create Calendar Event")
        if create_calendar:
            meeting_title = st.text_input("Meeting Title")
            meeting_date = st.date_input("Meeting Date")
            meeting_time = st.time_input("Meeting Time")

        # Add file size limit warning
        st.info("Maximum file size: 100MB")

    # Tabs for different input methods
    tab1, tab2 = st.tabs(["Upload Recording", "Meeting URL"])
    
    with tab1:
        uploaded_file = st.file_uploader(
            "Choose a meeting recording file", 
            type=['mp3', 'wav', 'm4a', 'ogg', 'mp4', 'avi', 'mov'],
            help="Supported formats: MP3, WAV, M4A, OGG, MP4, AVI, MOV"
        )

    with tab2:
        meeting_url = st.text_input(
            "Enter meeting recording URL",
            help="Supported platforms: Zoom, Microsoft Teams, Google Meet"
        )
        if meeting_url:
            if st.button("Download and Process"):
                with st.spinner("Downloading meeting recording..."):
                    file_content = download_meeting_recording(meeting_url)
                    if file_content:
                        # Determine file extension from URL
                        file_extension = os.path.splitext(urlparse(meeting_url).path)[1]
                        if not file_extension:
                            file_extension = '.mp4'  # Default to mp4
                        file_name = f"meeting_recording{file_extension}"
                        
                        # Process the downloaded recording
                        temp_file_path = process_meeting_recording(file_content, file_name)
                        if temp_file_path:
                            process_meeting_file(temp_file_path)

    if uploaded_file is not None:
        # Check file size (100MB limit)
        if uploaded_file.size > 100 * 1024 * 1024:  # 100MB in bytes
            st.error("File size exceeds 100MB limit. Please upload a smaller file.")
            return

        with st.spinner("Processing your file..."):
            try:
                # Save uploaded file temporarily
                with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as temp_file:
                    temp_file.write(uploaded_file.getvalue())
                    temp_file_path = temp_file.name

                process_meeting_file(temp_file_path)

            except Exception as e:
                st.error(f"An error occurred: {str(e)}")
                st.error("Please try again with a different file or contact support if the problem persists.")

if __name__ == "__main__":
    main() 
